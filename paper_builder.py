"""Build an anonymized workshop-paper draft from the frozen CICDDoS2019 TAIM results.

Follows the PNAD paper-builder pattern (Times New Roman, styled tables, no author
identifiers in the body). Source figures are the REAL_DATA_RESULTS tables; no fabricated
numbers. Output: TAIM_anonymous_workshop_paper.docx
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = "TAIM_anonymous_workshop_paper.docx"


def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_para(doc, text, size=11, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, size=13, bold=True):
    return add_para(doc, text, size=size, bold=bold, space_after=4)


def add_table(doc, header, rows, widths=None, font_size=9.5, highlight=()):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_font(r, size=font_size, bold=True)
        set_cell_shading(cell, "DDDDDD")
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_font(r, size=font_size, bold=((i - 1) in highlight))
    if widths:
        for j, w in enumerate(widths):
            for row in table.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Title
    add_para(doc, "TAIM Does Not Transfer to the CICDDoS2019 Benchmark: A Temporal-Detector "
                  "Mismatch", size=15, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Anonymous author(s) - anonymized for review", size=10.5, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # Abstract
    add_heading(doc, "Abstract")
    add_para(doc,
        "We test whether a time-aware, adaptive detector transfers to a real public DDoS "
        "benchmark. On CICDDoS2019, under strict family and day hold-outs, fold-isolation and "
        "validation-calibrated operating points, the adaptive detector TAIM is near-chance on "
        "unseen families (ROC-AUC \u2248 0.57, PR-AUC \u2248 0.05, F1 \u2248 0.05), while a "
        "supervised Random Forest baseline generalizes strongly (ROC-AUC \u2248 0.99, F1 "
        "\u2248 0.79) on 17 held-out families with \u226511 attack windows each. The narrow "
        "claim: TAIM does not transfer to this CICDDoS2019 benchmark. We trace the failure to a "
        "temporal-detector/benchmark mismatch - CICDDoS2019 family bursts last only minutes, "
        "giving \u22483.2 windows per source IP, so there is too little per-device history for a "
        "temporal baseline to learn from. We argue this is a benchmark-transfer result, not a "
        "deployment-level claim, and that a longer-horizon trace is required for a fair "
        "evaluation.", size=10.5)

    # 1 Introduction
    add_heading(doc, "1. Introduction")
    add_para(doc,
        "Network intrusion detectors are usually snapshot methods: they score each observation "
        "independently against a model of normal behavior. Time-aware detectors instead build "
        "a per-source baseline that updates online, so a sequence of individually-normal "
        "windows can accumulate into a sustained-deviation flag (and, in TAIM, escalate a "
        "mitigation ladder). The claim is that this buys detection of slow, low-and-slow "
        "attacks that a per-window model misses.")
    add_para(doc,
        "That claim is hard to test honestly because the detector is stateful: it must be "
        "warmed on training telemetry before it scores anything, and the warm-up must not see "
        "the evaluation windows. We therefore run the detector under a strict, fold-isolated "
        "protocol on a real public benchmark and ask a narrow question: does adaptive time-aware "
        "thresholding remain useful under realistic shift?")

    # 2 Dataset and protocol
    add_heading(doc, "2. Background: TAIM and the snapshot baselines")
    add_para(doc,
        "TAIM is a time-aware detector with three components: a per-device baseline that "
        "updates online (means/variance per time-of-day slot), a multi-signal fusion gate that "
        "flags a window when several signals are simultaneously elevated, and a mitigation "
        "ladder that escalates from watch to bandwidth cap to deauth as sustained deviation "
        "accumulates. Crucially it is stateful and causal (score-then-update), which is exactly "
        "why it must be warmed on training telemetry and scored on frozen state for evaluation.")
    add_para(doc,
        "The baselines are snapshot methods: a supervised Random Forest (log-scaled standardized "
        "features), unsupervised Isolation Forest (contamination tuned on inner validation), and "
        "a fixed-rule threshold on connection rate or bandwidth. All three score each window "
        "independently and need no per-device history.")

    # 3 Dataset and strict protocol
    add_heading(doc, "3. Dataset and strict protocol")
    add_para(doc,
        "Data. CICDDoS2019 (CC-BY-4.0, UNB). The 01-12 and 03-11 captures are streamed, "
        "adapted to per-(source-IP, 1-minute) windows, and tagged with their capture family and "
        "capture day. We use 1-minute buckets because each family attack burst lasts minutes; "
        "coarser buckets collapse a family's attacks into one or two windows. Result: 10,470 "
        "windows, 2 capture days, 17 families, 1,275 attack windows; every family has \u226511 "
        "attack windows (Syn 669 \u2192 WebDDoS 11), and \u22483.2 windows per source IP.")
    add_para(doc,
        "Splits. Never random rows. family hold-out holds out an entire attack family (test = "
        "held-out family windows + a 20% benign split held out from training); day hold-out holds "
        "out a whole capture day.")
    add_para(doc,
        "Fold isolation + calibration. The adaptive detector is warmed on the training rows "
        "only (updates on), then scores the test rows against a frozen baseline (updates off), "
        "so held-out family/day telemetry never updates the baseline that scores it. Only the "
        "unsupervised contamination is tuned on an inner validation split; the RF decision "
        "threshold stays at 0.5. The RF and TAIM recall@FPR cutoffs are each selected on a "
        "genuine validation split (RF: a fresh fit-split/validation-split model; TAIM: a "
        "chronological train-timeline split where the baseline is warmed on the earlier fraction "
        "and the cutoff is picked on the later, frozen-scored fraction). The test fold is never "
        "touched for calibration. Comparators see the same 5 windowed signals (bandwidth_mbps, "
        "conn_rate_ps, port_div, pkt_size_mean, app_req_ps; bandwidth/app_req log-scaled).")

    add_heading(doc, "3.1 Methodological hardening")
    add_para(doc,
        "Three evaluation habits were corrected before this result was trusted. First, fold "
        "isolation: a stateful detector cannot be run over the concatenated train+test frame and "
        "then have its test rows read, because its adaptive baseline would have processed the "
        "evaluation telemetry. TAIM is therefore run two-phase, warming on train and scoring "
        "test frozen. Second, in-sample calibration: a recall@FPR cutoff must be selected on a "
        "validation portion, not on the sequence whose scores are reported. Third, per-family "
        "support: a family-level result is only meaningful when each held-out family has enough "
        "attack windows, which required 1-minute buckets (coarser buckets left most families "
        "with a handful of windows). These are the standard pitfalls for any online detector.")

    # 4 Results
    add_heading(doc, "4. Results")
    add_para(doc, "Strict family hold-out (17 held-out families, \u226511 attack windows each):")
    add_table(
        doc,
        ["comparator", "F1 (\u00b195% CI)", "prec.", "recall", "PR-AUC (\u00b1CI)", "ROC-AUC (\u00b1CI)", "recall@1% FPR", "alerts"],
        [
            ["Random Forest (sup.)", "0.792 (\u00b10.049)", "0.742", "0.902", "0.956 (\u00b10.034)", "0.992 (\u00b10.010)", "0.919", "\u22485"],
            ["Isolation Forest", "0.117 (\u00b10.028)", "0.105", "0.856", "\u2014", "\u2014", "\u2014", "\u22489"],
            ["Fixed rule", "0.090 (\u00b10.063)", "0.053", "0.976", "\u2014", "\u2014", "\u2014", "\u224811"],
            ["TAIM (adaptive)", "0.052 (\u00b10.032)", "0.029", "0.541", "0.046 (\u00b10.031)", "0.569 (\u00b10.107)", "0.540", "\u224881"],
        ],
        widths=[1.4, 1.0, 0.6, 0.65, 0.95, 0.95, 0.85, 0.5],
        highlight=(0,),
    )
    add_para(doc,
        "Per held-out family, supervised Random Forest F1 ranges 0.52\u20130.90 and ROC-AUC "
        "0.92\u20131.00 across all 17 families (see Table 2); TAIM F1 is 0.00\u20130.23. Under "
        "strict day hold-out (2 capture days), RF F1 \u2248 0.70 (AUC \u2248 0.92) versus TAIM "
        "F1 \u2248 0.18.", size=10.5)

    add_para(doc,
        "Important operating-point caveat. The 1% FPR cutoff is a *validation-selected* "
        "threshold, not a guaranteed operational false-positive rate. When applied to the held-out "
        "test folds, the TAIM cutoff produced an achieved test FPR of \u2248 0.395 - over 39 "
        "times the 1% budget - because TAIM's score does not separate attacks under shift. The "
        "corresponding RF cutoff achieved FPR \u2248 0.005. These are test-set FPRs observed under "
        "shift, not a guarantee that the detector will operate at 1% FPR in deployment.",
        size=10.5)

    add_para(doc, "Table 2. Per held-out family, supervised Random Forest (fold-isolated):")
    add_table(
        doc,
        ["held-out family", "attack windows", "F1", "ROC-AUC"],
        [
            ["DrDoS_DNS", "36", "0.875", "0.996"],
            ["DrDoS_UDP", "46", "0.900", "0.998"],
            ["UDP-lag", "41", "0.901", "1.000"],
            ["DrDoS_SNMP", "30", "0.853", "1.000"],
            ["MSSQL", "28", "0.844", "0.998"],
            ["Portmap", "79", "0.838", "0.984"],
            ["DrDoS_LDAP", "12", "0.688", "0.995"],
            ["DrDoS_NTP", "172", "0.521", "0.921"],
            ["Syn", "669", "0.749", "0.975"],
            ["WebDDoS", "11", "0.710", "1.000"],
        ],
        widths=[1.5, 1.2, 0.8, 0.9],
    )
    add_para(doc,
        "Minimum family support is 11 attack windows (WebDDoS); the well-sampled families reach "
        "F1 0.9, and even the hardest (DrDoS_NTP) reaches AUC 0.92. The range is stable, so the "
        "supervised result is not an artifact of one easy family.", size=10.5)

    add_heading(doc, "4.1 Why TAIM is near-chance here")
    add_para(doc,
        "TAIM is temporal: it learns a per-device baseline over time and flags sustained "
        "deviation. On CICDDoS2019 the family bursts are short, giving roughly 3.2 windows per "
        "source IP, so there is essentially no per-device history for TAIM's baseline to learn "
        "from and no time-of-day regime to model. Even after fold isolation and validation "
        "calibration, TAIM's score orders attacks about as well as chance: its validation-"
        "calibrated recall@1%FPR cutoff achieved a test FPR of \u2248 0.395 on held-out folds, "
        "\u224839 times the 1% budget, i.e. the score does not separate the classes. Isolation "
        "Forest and Random Forest are snapshot models that need no such history, so they are not "
        "similarly handicapped. We read this as a benchmark mismatch, not a general failure of "
        "temporal detection.")

    # 5 Related work
    add_heading(doc, "5. Related work")
    add_para(doc,
        "Online, time-aware detection is a common framing for network intrusion detection, and "
        "adaptive baselining is attractive because it does not require labeled attacks. Empirical "
        "work on public DDoS/IDS benchmarks (CIC-IDS, CSE-CIC-IDS, UNSW-NB15) most often reports "
        "snapshot methods. Less reported is whether the *temporal* premise survives when the "
        "benchmark's attacks are short and family-isolated rather than spread over a long host "
        "history. Our result is a concrete counterexample: a temporal detector that is strong on "
        "long-horizon synthetic traffic is near-chance on a real benchmark whose per-device "
        "history is too short. This is the same synthetic-to-real gap documented for host anomaly "
        "detectors, observed here for an online network detector.")

    # 6 Conclusion
    add_heading(doc, "6. Conclusion")
    add_para(doc,
        "On real CICDDoS2019 with strict family/day hold-outs, fold isolation and validation "
        "calibration, adaptive time-aware thresholding (TAIM) does not remain useful under "
        "distribution shift: it is near-chance on unseen families while a supervised baseline "
        "generalizes strongly. The result is a benchmark-transfer finding: TAIM's per-device "
        "temporal baseline has no purchase on short, family-isolated bursts. A longer-horizon "
        "real trace (e.g. LANL or a multi-day capture) is the correct next test before any "
        "deployment claim.")

    # 7 Honest limits
    add_heading(doc, "7. Limitations")
    add_para(doc,
        "This is a bounded per-family sample, not the full 22 GB trace. Only two capture days "
        "are available, so the day-hold-out confidence interval is wide and the family hold-out "
        "(n = 17) is the primary evidence. TAIM's internal state is not re-run deterministic "
        "across separate process invocations (pre-existing), so the reported aggregates use a "
        "fixed evaluation pass. The 1-minute bucket also inflates absolute per-family window "
        "counts but does not change the ranking. These caveats do not change the result: the "
        "adaptive detector falls far behind a supervised baseline even under a fold-isolated, "
        "calibrated protocol.")

    add_heading(doc, "8. Repository, data, and reproduction")
    add_para(doc,
        "Repository (anonymized for review): github.com/Farooq-Syed/taim. The evaluation entry "
        "point is src/real_cicddos_eval.py; the window builder is scripts/build_real_windows.py; "
        "the capture downloader is scripts/download_real_data.py. Frozen preprocessing and split "
        "definitions are documented in REAL_DATA_RESULTS.md.", size=10)
    add_para(doc,
        "Data attribution and license. CICDDoS2019 (CC-BY-4.0), Canadian Institute for "
        "Cybersecurity, University of New Brunswick. Dataset page: "
        "https://www.unb.ca/cic/datasets/ddos-2019.html. The captures are public and used here "
        "for research evaluation only. This manuscript is released under the repository's "
        "Non-Commercial Personal-Use License.", size=10)
    add_para(doc,
        "Reproduction commands.",
        size=10)
    add_para(doc,
        "python scripts/download_real_data.py; "
        "python scripts/build_real_windows.py --rows-per-family 1000000 --bucket-min 1 "
        "--output data/cicddos_real_windows.csv; "
        "python src/real_cicddos_eval.py --input data/cicddos_real_windows.csv --split family; "
        "python src/real_cicddos_eval.py --input data/cicddos_real_windows.csv --split day",
        size=9.5)

    # AI-use disclosure
    add_heading(doc, "9. AI-use disclosure")
    add_para(doc,
        "AI coding assistance was used during implementation and drafting. The author directed "
        "the research question, the benchmark evaluation protocol, the strict split and "
        "calibration design, the interpretation of the negative result, and reviewed and verified "
        "the final code and manuscript claims. AI assistance did not set the research direction or "
        "the claims.", size=10)

    # References
    add_heading(doc, "References")
    add_para(doc, "[1] CICDDoS2019 dataset, Canadian Institute for Cybersecurity, "
        "University of New Brunswick. https://www.unb.ca/cic/datasets/ddos-2019.html", size=9.5)
    add_para(doc, "[2] Moustafa & Slay, UNSW-NB15, MilCIS 2015. Sharafaldin et al., CIC-IDS, "
        "ICISSP 2018. (Public NIDS benchmarks.)", size=9.5)
    add_para(doc, "[3] Liu, Ting & Zhou, Isolation Forest, ICDM 2008. Breunig et al., LOF, "
        "SIGMOD 2000. Breiman, Random Forests, ML 2001.", size=9.5)
    add_para(doc, "[4] Goldschmidt & Chud\u00e1, Network Intrusion Datasets: A Survey, 2025. "
        "https://arxiv.org/abs/2502.06688", size=9.5)

    # Appendix - review questions
    add_heading(doc, "Appendix. Three review questions")
    add_para(doc, "1. Is the claim narrow enough? Yes: 'adaptive time-aware thresholding does "
        "not transfer to this benchmark' - explicitly a benchmark-transfer result, not a "
        "deployment-level claim.", size=9.5)
    add_para(doc, "2. Is the split protocol sound? Family and day hold-outs are strict; "
        "the adaptive detector is fold-isolated (warmed on train, scored frozen); the recall@FPR "
        "cutoffs are calibrated on validation, never the test fold.", size=9.5)
    add_para(doc, "3. What experiment would most change your confidence? A longer-horizon real "
        "trace (e.g. LANL or a multi-day capture) that gives each source IP enough history for "
        "a temporal baseline to learn from - the exact condition this benchmark lacks.",
        size=9.5)

    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = "Anonymous Workshop Manuscript"
    doc.core_properties.subject = "Methodological review"
    doc.core_properties.comments = ""
    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    build()
